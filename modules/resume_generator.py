from typing import Dict, List
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.graphics.shapes import Line, Drawing
from reportlab.platypus import KeepTogether
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

class ResumeGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom styles for the resume"""
        # Name style - large, bold, centered
        self.name_style = ParagraphStyle(
            'NameStyle',
            parent=self.styles['Heading1'],
            fontSize=20,
            spaceAfter=6,
            spaceBefore=0,
            alignment=1,  # Center alignment
            textColor=colors.black,
            fontName='Times-Bold',
            leading=24
        )
        
        # Contact info style - smaller, centered
        self.contact_style = ParagraphStyle(
            'ContactInfo',
            parent=self.styles['Normal'],
            fontSize=10,
            alignment=1,  # Center alignment
            spaceAfter=8,
            spaceBefore=0,
            fontName='Times-Roman',
            textColor=colors.black,
            leading=12
        )
        
        # Section header style - 11pt as requested
        self.section_style = ParagraphStyle(
            'SectionHeader',
            parent=self.styles['Heading2'],
            fontSize=11,
            spaceAfter=3,
            spaceBefore=8,
            textColor=colors.black,
            fontName='Times-Bold',
            leftIndent=0,
            leading=13
        )
        
        # Body text style - 10pt as requested
        self.body_style = ParagraphStyle(
            'CustomBody',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=3,
            spaceBefore=0,
            fontName='Times-Roman',
            textColor=colors.black,
            leading=11,
            leftIndent=0
        )
        
        # Job title style - bold for position titles
        self.job_title_style = ParagraphStyle(
            'JobTitle',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=1,
            spaceBefore=3,
            fontName='Times-Bold',
            textColor=colors.black,
            leading=11
        )
        
        # Company/Institution style - italic
        self.company_style = ParagraphStyle(
            'CompanyStyle',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=1,
            spaceBefore=0,
            fontName='Times-Italic',
            textColor=colors.black,
            leading=11
        )
        
        # Duration style - right aligned, smaller
        self.duration_style = ParagraphStyle(
            'DurationStyle',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=1,
            spaceBefore=0,
            fontName='Times-Roman',
            textColor=colors.black,
            alignment=2,  # Right alignment
            leading=11
        )
        
        # Description style - slightly indented
        self.description_style = ParagraphStyle(
            'DescriptionStyle',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=2,
            spaceBefore=1,
            fontName='Times-Roman',
            textColor=colors.black,
            leading=11,
            leftIndent=10,
            bulletIndent=0
        )
        
        # Skills category style - for subsection titles
        self.skills_category_style = ParagraphStyle(
            'SkillsCategoryStyle',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=1,
            spaceBefore=2,
            fontName='Times-Bold',
            textColor=colors.black,
            leading=11,
            leftIndent=0
        )
        
        # Skills content style - for skills list
        self.skills_content_style = ParagraphStyle(
            'SkillsContentStyle',
            parent=self.styles['Normal'],
            fontSize=10,
            spaceAfter=1,
            spaceBefore=0,
            fontName='Times-Roman',
            textColor=colors.black,
            leading=11,
            leftIndent=0
        )
    
    def create_section_line(self) -> Drawing:
        """Create a horizontal line for section separation"""
        drawing = Drawing(7.5*inch, 0.05*inch)  # Width adjusted for smaller margins
        line = Line(0, 0.025*inch, 7.5*inch, 0.025*inch)
        line.strokeColor = colors.black
        line.strokeWidth = 0.5
        drawing.add(line)
        return drawing
    
    def prioritize_experiences(self, experiences: List[Dict], job_analysis: Dict) -> List[Dict]:
        """Prioritize and tailor experiences based on job requirements"""
        if not job_analysis:
            return experiences
        
        # Get job keywords and skills
        job_keywords = [kw[0].lower() for kw in job_analysis.get('keywords', [])]
        job_skills = [skill.lower() for skill in 
                     job_analysis.get('skills', {}).get('technical', []) + 
                     job_analysis.get('skills', {}).get('soft', [])]
        
        all_job_terms = set(job_keywords + job_skills)
        
        # Score each experience based on relevance
        scored_experiences = []
        for exp in experiences:
            score = 0
            exp_text = (exp.get('title', '') + ' ' + 
                       exp.get('company', '') + ' ' + 
                       exp.get('description', '')).lower()
            
            # Count matches with job terms
            for term in all_job_terms:
                if term in exp_text:
                    score += 1
            
            # Boost score for priority skills
            priority_skills = job_analysis.get('priority_skills', [])
            for priority in priority_skills[:5]:  # Top 5 priority skills
                if isinstance(priority, dict):
                    skill_name = priority.get('skill', '')
                    if skill_name and skill_name.lower() in exp_text:
                        # Check for in_requirements (regular analyzer) or importance (LLM analyzer)
                        is_high_priority = priority.get('in_requirements', False) or priority.get('importance') == 'high'
                        score += 3 if is_high_priority else 2
            
            scored_experiences.append((score, exp))
        
        # Sort by score (descending) and return experiences
        scored_experiences.sort(key=lambda x: x[0], reverse=True)
        return [exp for score, exp in scored_experiences]
    
    def tailor_experience_description(self, description: str, job_analysis: Dict) -> str:
        """Enhance experience description with job-relevant keywords"""
        if not job_analysis or not description:
            return description
        
        # Get priority skills and keywords
        priority_skills = job_analysis.get('priority_skills', [])
        top_keywords = [kw[0] for kw in job_analysis.get('keywords', [])[:10]]
        
        # For now, return original description
        # In a more advanced version, this could use NLP to enhance descriptions
        return description
    
    def get_all_skills_combined(self, profile: Dict) -> str:
        """Combine all skill categories into a single string"""
        all_skills = []
        
        # Collect skills from all categories
        categories = ['programming_skills', 'technologies', 'language_skills', 'certifications']
        for category in categories:
            skills_text = profile.get(category, '')
            if skills_text:
                skills_list = [skill.strip() for skill in skills_text.split(',') if skill.strip()]
                all_skills.extend(skills_list)
        
        return ', '.join(all_skills)
    
    def generate_tailored_skills(self, profile: Dict, job_analysis: Dict) -> str:
        """Generate a tailored skills section based on job requirements"""
        # Get all skills combined from categories
        all_skills_text = self.get_all_skills_combined(profile)
        
        if not all_skills_text:
            return ""
        
        profile_skill_list = [skill.strip() for skill in all_skills_text.split(',')]
        
        if not job_analysis:
            return all_skills_text
        
        # Get job skills
        job_technical = job_analysis.get('skills', {}).get('technical', [])
        job_soft = job_analysis.get('skills', {}).get('soft', [])
        # Handle different priority skill formats
        priority_skills = []
        for ps in job_analysis.get('priority_skills', []):
            if isinstance(ps, dict):
                skill_name = ps.get('skill', '')
                if skill_name:
                    priority_skills.append(skill_name)
            else:
                priority_skills.append(str(ps))
        
        # Prioritize skills that match the job
        matched_skills = []
        unmatched_skills = []
        
        for skill in profile_skill_list:
            skill_lower = skill.lower()
            is_match = (skill_lower in [js.lower() for js in job_technical + job_soft] or
                       skill in priority_skills)
            
            if is_match:
                matched_skills.append(skill)
            else:
                unmatched_skills.append(skill)
        
        # Put matched skills first
        tailored_skills = matched_skills + unmatched_skills
        return ', '.join(tailored_skills)
    
    def generate_categorized_skills_text(self, profile: Dict) -> str:
        """Generate skills text organized by categories"""
        skills_sections = []
        
        categories = {
            'Programming': profile.get('programming_skills', ''),
            'Technologies': profile.get('technologies', ''),
            'Languages': profile.get('language_skills', ''),
            'Certifications': profile.get('certifications', '')
        }
        
        for category, skills_text in categories.items():
            if skills_text:
                skills_sections.append(f"{category}: {skills_text}")
        
        return ' | '.join(skills_sections)
    
    def generate_categorized_skills_for_pdf(self, profile: Dict) -> Dict[str, str]:
        """Generate skills organized by categories for PDF formatting"""
        categories = {
            'Programming': profile.get('programming_skills', ''),
            'Technologies': profile.get('technologies', ''),
            'Languages': profile.get('language_skills', ''),
            'Certifications': profile.get('certifications', '')
        }
        
        # Return only categories that have content
        return {category: skills for category, skills in categories.items() if skills.strip()}
    
    def enhance_content_with_llm(self, content: str, content_type: str, job_analysis: Dict = None) -> str:
        """
        Placeholder for future LLM integration to enhance resume content.
        
        Args:
            content: The original content to enhance
            content_type: Type of content ('summary', 'experience', 'skills', etc.)
            job_analysis: Job analysis data for context
            
        Returns:
            Enhanced content (currently returns original content)
        """
        # TODO: Integrate with LLM API (OpenAI, Claude, etc.) to:
        # 1. Enhance professional summary for better impact
        # 2. Improve experience descriptions with action verbs and quantifiable results
        # 3. Optimize skills presentation for ATS compatibility
        # 4. Tailor content specifically for the job requirements
        # 5. Fix grammar and enhance language quality
        
        # For now, return original content
        # Future implementation might look like:
        # if content_type == 'summary':
        #     return self._enhance_summary_with_llm(content, job_analysis)
        # elif content_type == 'experience':
        #     return self._enhance_experience_with_llm(content, job_analysis)
        # etc.
        
        return content
    
    def generate_professional_summary(self, profile: Dict, job_analysis: Dict) -> str:
        """Generate a tailored professional summary"""
        base_summary = profile.get('summary', '')
        
        if not job_analysis:
            return base_summary or "Experienced professional with a strong background in technology and innovation."
        
        # Extract key elements from job analysis
        experience_level = job_analysis.get('experience_level', 'experienced')
        top_skills = job_analysis.get('priority_skills', [])[:3]
        
        if base_summary:
            return base_summary
        
        # Generate a basic summary if none exists
        skill_text = ""
        if top_skills:
            skills_list = [skill['skill'] for skill in top_skills]
            skill_text = f" with expertise in {', '.join(skills_list)}"
        
        level_text = {
            'entry': 'Motivated entry-level professional',
            'mid': 'Experienced professional',
            'senior': 'Senior professional with proven leadership',
            'executive': 'Executive-level professional'
        }.get(experience_level, 'Experienced professional')
        
        return f"{level_text}{skill_text}, ready to contribute to organizational success."
    
    def generate_pdf_resume(self, profile: Dict, job_analysis: Dict = None) -> io.BytesIO:
        """Generate a professional PDF resume"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter, 
            topMargin=0.75*inch,
            bottomMargin=0.75*inch,
            leftMargin=0.5*inch,  # Reduced from 0.75 to 0.5
            rightMargin=0.5*inch  # Reduced from 0.75 to 0.5
        )
        content = []
        
        # Header Section - Name
        name = profile.get('name', 'Your Name')
        content.append(Paragraph(name, self.name_style))
        
        # Contact information - formatted as a single line
        contact_parts = []
        if profile.get('email'):
            contact_parts.append(profile['email'])
        if profile.get('phone'):
            contact_parts.append(profile['phone'])
        if profile.get('location'):
            contact_parts.append(profile['location'])
        if profile.get('linkedin'):
            contact_parts.append(profile['linkedin'])
        
        if contact_parts:
            contact_text = ' • '.join(contact_parts)
            content.append(Paragraph(contact_text, self.contact_style))
        
        # Professional Summary
        summary = self.generate_professional_summary(profile, job_analysis)
        if summary:
            # Apply LLM enhancement placeholder
            enhanced_summary = self.enhance_content_with_llm(summary, 'summary', job_analysis)
            content.append(Paragraph("PROFESSIONAL SUMMARY", self.section_style))
            content.append(self.create_section_line())
            content.append(Paragraph(enhanced_summary, self.body_style))
        
        # Core Competencies Section - Categorized
        skills_categories = self.generate_categorized_skills_for_pdf(profile)
        if skills_categories:
            content.append(Paragraph("CORE COMPETENCIES", self.section_style))
            content.append(self.create_section_line())
            
            # If tailored, prioritize relevant skills within each category
            if job_analysis:
                # For tailored resumes, show most relevant categories first
                job_skills = job_analysis.get('skills', {}).get('technical', []) + job_analysis.get('skills', {}).get('soft', [])
                job_skills_lower = [skill.lower() for skill in job_skills]
                
                # Score categories by relevance
                category_scores = {}
                for category, skills_text in skills_categories.items():
                    score = 0
                    skills_list = [s.strip().lower() for s in skills_text.split(',') if s.strip()]
                    for skill in skills_list:
                        if skill in job_skills_lower:
                            score += 1
                    category_scores[category] = score
                
                # Sort categories by relevance score
                sorted_categories = sorted(skills_categories.items(), key=lambda x: category_scores.get(x[0], 0), reverse=True)
            else:
                # For standard resumes, use default order
                sorted_categories = skills_categories.items()
            
            # Display each skills category
            for category, skills_text in sorted_categories:
                if skills_text.strip():
                    # Category title (bold)
                    content.append(Paragraph(f"{category}:", self.skills_category_style))
                    
                    # Skills content
                    enhanced_skills = self.enhance_content_with_llm(skills_text, 'skills', job_analysis)
                    content.append(Paragraph(enhanced_skills, self.skills_content_style))
        
        # Professional Experience
        experiences = profile.get('experiences', [])
        if experiences:
            content.append(Paragraph("PROFESSIONAL EXPERIENCE", self.section_style))
            content.append(self.create_section_line())
            
            # Prioritize experiences based on job analysis
            prioritized_exp = self.prioritize_experiences(experiences, job_analysis)
            
            for i, exp in enumerate(prioritized_exp):
                # Create a table for each experience entry for better layout
                exp_data = []
                
                # Row 1: Job Title and Duration
                title = exp.get('title', 'Position')
                duration = exp.get('duration', '')
                
                title_cell = Paragraph(title, self.job_title_style)
                duration_cell = Paragraph(duration, self.duration_style)
                exp_data.append([title_cell, duration_cell])
                
                # Row 2: Company
                if exp.get('company'):
                    company_cell = Paragraph(exp.get('company'), self.company_style)
                    empty_cell = Paragraph("", self.company_style)
                    exp_data.append([company_cell, empty_cell])
                
                # Create table for job header (adjusted for smaller margins)
                exp_table = Table(exp_data, colWidths=[5.5*inch, 2*inch])
                exp_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
                ]))
                
                content.append(exp_table)
                
                # Description with bullet points
                description = self.tailor_experience_description(
                    exp.get('description', ''), job_analysis
                )
                if description:
                    # Apply LLM enhancement placeholder
                    enhanced_description = self.enhance_content_with_llm(description, 'experience', job_analysis)
                    
                    # Split description into bullet points if it contains multiple sentences
                    sentences = [s.strip() for s in enhanced_description.split('.') if s.strip()]
                    if len(sentences) > 1:
                        for sentence in sentences:
                            if sentence:
                                bullet_text = f"• {sentence}."
                                content.append(Paragraph(bullet_text, self.description_style))
                    else:
                        bullet_text = f"• {enhanced_description}"
                        content.append(Paragraph(bullet_text, self.description_style))
                
                # Add space between experiences (except for the last one)
                if i < len(prioritized_exp) - 1:
                    content.append(Spacer(1, 0.08*inch))
        
        # Education
        education = profile.get('education', [])
        if education:
            content.append(Paragraph("EDUCATION", self.section_style))
            content.append(self.create_section_line())
            
            for i, edu in enumerate(education):
                # Create table for education entry
                edu_data = []
                
                # Degree and Field
                degree_text = edu.get('degree', 'Degree')
                if edu.get('field'):
                    degree_text += f" in {edu.get('field')}"
                
                year = edu.get('year', '')
                
                degree_cell = Paragraph(degree_text, self.job_title_style)
                year_cell = Paragraph(year, self.duration_style)
                edu_data.append([degree_cell, year_cell])
                
                # Institution
                if edu.get('institution'):
                    institution_cell = Paragraph(edu.get('institution'), self.company_style)
                    empty_cell = Paragraph("", self.company_style)
                    edu_data.append([institution_cell, empty_cell])
                
                # Create table
                edu_table = Table(edu_data, colWidths=[4.5*inch, 2*inch])
                edu_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 0),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                    ('TOPPADDING', (0, 0), (-1, -1), 0),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
                ]))
                
                content.append(edu_table)
                
                # Additional details
                if edu.get('details'):
                    details_text = f"• {edu.get('details')}"
                    content.append(Paragraph(details_text, self.description_style))
                
                # Add space between education entries (except for the last one)
                if i < len(education) - 1:
                    content.append(Spacer(1, 0.05*inch))
        
        # Build PDF with better page breaking
        doc.build(content)
        buffer.seek(0)
        return buffer
    
    def generate_word_resume(self, profile: Dict, job_analysis: Dict = None) -> io.BytesIO:
        """Generate a professional Word document resume"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Header - Name (large, centered, bold)
        name = profile.get('name', 'Your Name')
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.font.size = Inches(0.3)  # Large font
        name_run.font.bold = True
        name_run.font.name = 'Calibri'
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact information (centered, smaller)
        contact_info = []
        if profile.get('email'):
            contact_info.append(profile['email'])
        if profile.get('phone'):
            contact_info.append(profile['phone'])
        if profile.get('location'):
            contact_info.append(profile['location'])
        if profile.get('linkedin'):
            contact_info.append(profile['linkedin'])
        
        if contact_info:
            contact_para = doc.add_paragraph(' • '.join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.runs[0]
            contact_run.font.size = Inches(0.12)
            contact_run.font.name = 'Calibri'
        
        # Add space after header
        doc.add_paragraph()
        
        # Professional Summary
        summary = self.generate_professional_summary(profile, job_analysis)
        if summary:
            summary_heading = doc.add_heading('PROFESSIONAL SUMMARY', level=2)
            summary_heading.runs[0].font.name = 'Calibri'
            summary_heading.runs[0].font.size = Inches(0.16)
            
            summary_para = doc.add_paragraph(summary)
            summary_para.runs[0].font.name = 'Calibri'
            summary_para.runs[0].font.size = Inches(0.12)
        
        # Skills
        if job_analysis:
            skills = self.generate_tailored_skills(profile, job_analysis)
            skills_heading = doc.add_heading('CORE COMPETENCIES', level=2)
        else:
            skills = self.generate_categorized_skills_text(profile)
            skills_heading = doc.add_heading('TECHNICAL SKILLS', level=2)
        
        if skills:
            skills_heading.runs[0].font.name = 'Calibri'
            skills_heading.runs[0].font.size = Inches(0.16)
            
            # Format skills professionally
            skills_formatted = skills.replace(' | ', ' • ').replace(': ', ': ')
            skills_para = doc.add_paragraph(skills_formatted)
            skills_para.runs[0].font.name = 'Calibri'
            skills_para.runs[0].font.size = Inches(0.12)
        
        # Professional Experience
        experiences = profile.get('experiences', [])
        if experiences:
            exp_heading = doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
            exp_heading.runs[0].font.name = 'Calibri'
            exp_heading.runs[0].font.size = Inches(0.16)
            
            # Prioritize experiences based on job analysis
            prioritized_exp = self.prioritize_experiences(experiences, job_analysis)
            
            for exp in prioritized_exp:
                # Job title (bold, larger)
                title = exp.get('title', 'Position')
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(title)
                title_run.font.bold = True
                title_run.font.name = 'Calibri'
                title_run.font.size = Inches(0.14)
                
                # Duration (right-aligned on same line if possible)
                if exp.get('duration'):
                    duration_run = title_para.add_run(f" — {exp.get('duration')}")
                    duration_run.font.name = 'Calibri'
                    duration_run.font.size = Inches(0.11)
                    duration_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # Company (italic, smaller)
                if exp.get('company'):
                    company_para = doc.add_paragraph(exp.get('company'))
                    company_run = company_para.runs[0]
                    company_run.font.italic = True
                    company_run.font.name = 'Calibri'
                    company_run.font.size = Inches(0.12)
                    company_run.font.color.rgb = RGBColor(96, 96, 96)
                
                # Description with bullet points
                description = self.tailor_experience_description(
                    exp.get('description', ''), job_analysis
                )
                if description:
                    # Split into sentences and create bullet points
                    sentences = [s.strip() for s in description.split('.') if s.strip()]
                    if len(sentences) > 1:
                        for sentence in sentences:
                            if sentence:
                                bullet_para = doc.add_paragraph(f"• {sentence}.")
                                bullet_para.runs[0].font.name = 'Calibri'
                                bullet_para.runs[0].font.size = Inches(0.11)
                                bullet_para.paragraph_format.left_indent = Inches(0.2)
                    else:
                        bullet_para = doc.add_paragraph(f"• {description}")
                        bullet_para.runs[0].font.name = 'Calibri'
                        bullet_para.runs[0].font.size = Inches(0.11)
                        bullet_para.paragraph_format.left_indent = Inches(0.2)
                
                # Add space between experiences
                doc.add_paragraph()
        
        # Education
        education = profile.get('education', [])
        if education:
            edu_heading = doc.add_heading('EDUCATION', level=2)
            edu_heading.runs[0].font.name = 'Calibri'
            edu_heading.runs[0].font.size = Inches(0.16)
            
            for edu in education:
                # Degree and field (bold)
                degree_text = edu.get('degree', 'Degree')
                if edu.get('field'):
                    degree_text += f" in {edu.get('field')}"
                
                degree_para = doc.add_paragraph()
                degree_run = degree_para.add_run(degree_text)
                degree_run.font.bold = True
                degree_run.font.name = 'Calibri'
                degree_run.font.size = Inches(0.14)
                
                # Year
                if edu.get('year'):
                    year_run = degree_para.add_run(f" — {edu.get('year')}")
                    year_run.font.name = 'Calibri'
                    year_run.font.size = Inches(0.11)
                    year_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # Institution (italic)
                if edu.get('institution'):
                    institution_para = doc.add_paragraph(edu.get('institution'))
                    institution_run = institution_para.runs[0]
                    institution_run.font.italic = True
                    institution_run.font.name = 'Calibri'
                    institution_run.font.size = Inches(0.12)
                    institution_run.font.color.rgb = RGBColor(96, 96, 96)
                
                # Additional details
                if edu.get('details'):
                    details_para = doc.add_paragraph(f"• {edu.get('details')}")
                    details_para.runs[0].font.name = 'Calibri'
                    details_para.runs[0].font.size = Inches(0.11)
                    details_para.paragraph_format.left_indent = Inches(0.2)
                
                # Add space between education entries
                doc.add_paragraph()
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer